import io

from docx import Document as DocxDocument  # type: ignore[import-untyped]

from .base import ProcessingResult


def extract_docx(data: bytes) -> ProcessingResult:
    doc = DocxDocument(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return ProcessingResult(
        text="\n".join(parts),
        method="direct",
        page_count=None,
    )
